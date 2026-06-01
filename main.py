"""
FastAPI Backend for Offline File Chat with Local LLM Integration
Production-Grade Implementation
"""

import os
import json
import uuid
import asyncio
import re
from collections import Counter
from typing import Optional, Dict, List, Any
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import httpx
import uvicorn

# File extraction libraries
from pypdf import PdfReader
from docx import Document
import csv
import io

# ============================================================================
# CONFIGURATION
# ============================================================================

LOCAL_LLM_URL = "http://localhost:8080"
LOCAL_LLM_ENDPOINT = f"{LOCAL_LLM_URL}/v1/chat/completions"
UPLOAD_DIR = Path("./uploaded_files")
CONTEXT_DIR = Path("./file_contexts")

# Create directories if they don't exist
UPLOAD_DIR.mkdir(exist_ok=True)
CONTEXT_DIR.mkdir(exist_ok=True)

# ============================================================================
# GLOBAL STATE
# ============================================================================

app = FastAPI(
    title="Offline File Chat API",
    description="Local LLM + File Extraction",
    version="1.0.0"
)

# In-memory storage: session_id -> {file_id: text_content, ...}
file_contexts: Dict[str, Dict[str, str]] = {}

# In-memory conversation history: session_id -> list of messages
conversation_history: Dict[str, List[Dict]] = {}

# ============================================================================
# CORS MIDDLEWARE
# ============================================================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Local-only, safe for offline development
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================================
# REQUEST/RESPONSE MODELS
# ============================================================================

class ChatRequest(BaseModel):
    """Chat message request"""
    session_id: str
    message: str
    file_ids: Optional[List[str]] = None


class ChatResponse(BaseModel):
    """Chat response metadata"""
    session_id: str
    message_id: str
    status: str


class FileUploadResponse(BaseModel):
    """File upload response"""
    file_id: str
    filename: str
    size: int
    status: str
    extracted_chars: int
    extracted_text: str
    keyword_analysis: Optional[Dict[str, Any]] = None
    document_analysis: Optional[Dict[str, Any]] = None


# ============================================================================
# FILE EXTRACTION FUNCTIONS
# ============================================================================

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF using pypdf"""
    try:
        pdf_reader = PdfReader(io.BytesIO(file_content))
        text_content = ""
        for page_num, page in enumerate(pdf_reader.pages):
            text_content += f"\n--- Page {page_num + 1} ---\n"
            text_content += page.extract_text()
        return text_content.strip()
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        doc = Document(io.BytesIO(file_content))
        text_content = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text_content += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content += cell.text + " | "
                text_content += "\n"
        return text_content.strip()
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {str(e)}")


def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV"""
    try:
        decoded = file_content.decode("utf-8")
        reader = csv.DictReader(io.StringIO(decoded))
        text_content = ""
        for idx, row in enumerate(reader, 1):
            text_content += f"\nRow {idx}: {json.dumps(row)}\n"
        return text_content.strip()
    except Exception as e:
        raise ValueError(f"CSV extraction failed: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT"""
    try:
        return file_content.decode("utf-8")
    except Exception as e:
        raise ValueError(f"TXT extraction failed: {str(e)}")


async def extract_text_from_file(filename: str, file_content: bytes) -> str:
    """
    Route file extraction based on extension
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_content)
    elif filename_lower.endswith(".docx") or filename_lower.endswith(".doc"):
        return extract_text_from_docx(file_content)
    elif filename_lower.endswith(".csv"):
        return extract_text_from_csv(file_content)
    elif filename_lower.endswith(".txt"):
        return extract_text_from_txt(file_content)
    else:
        # Default: try to decode as text
        try:
            return extract_text_from_txt(file_content)
        except:
            raise ValueError(f"Unsupported file type: {filename_lower}")


def extract_keywords_from_text(text: str) -> Dict[str, Any]:
    """Extract important keywords and frequency from text."""
    words = re.findall(r"\b[a-zA-Z]{3,}\b", text.lower())
    stop_words = {
        "the", "and", "for", "with", "this", "that", "from", "are", "was", "were",
        "have", "has", "had", "into", "about", "your", "their", "been", "will", "shall",
        "can", "could", "would", "should", "there", "here", "when", "where", "which", "while",
        "page"
    }
    filtered_words = [word for word in words if word not in stop_words]
    keyword_counts = Counter(filtered_words).most_common(15)
    keywords = [{"word": word, "frequency": frequency} for word, frequency in keyword_counts]
    return {
        "keywords": keywords,
        "keyword_count": len(keywords)
    }


def analyze_document_structure(text: str) -> Dict[str, Any]:
    """Analyze pages, paragraphs, sentences and sentence length."""
    page_markers = re.findall(r"--- Page \d+ ---", text)
    page_count = len(page_markers) if page_markers else (1 if text.strip() else 0)

    stripped_text = text.strip()
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", stripped_text) if p.strip()]
    paragraph_count = len(paragraphs)

    sentence_candidates = [s.strip() for s in re.split(r"[.!?]+", stripped_text) if s.strip()]
    sentence_count = len(sentence_candidates)

    if sentence_count > 0:
        total_words = sum(len(sentence.split()) for sentence in sentence_candidates)
        avg_sentence_length = round(total_words / sentence_count, 2)
    else:
        avg_sentence_length = 0.0

    return {
        "page_count": page_count,
        "paragraph_count": paragraph_count,
        "sentence_count": sentence_count,
        "avg_sentence_length": avg_sentence_length,
    }


# ============================================================================
# ENDPOINTS
# ============================================================================

@app.post("/upload", response_model=FileUploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload and parse a document file.
    Supports: PDF, DOCX, CSV, TXT and other text-based formats.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        extracted_text = await extract_text_from_file(file.filename, content)
        is_pdf = file.filename.lower().endswith(".pdf")
        keyword_analysis = extract_keywords_from_text(extracted_text) if is_pdf else None
        document_analysis = analyze_document_structure(extracted_text) if is_pdf else None
        
        # Generate unique file ID
        file_id = str(uuid.uuid4())
        
        # Store extracted text in memory (indexed by file_id)
        # Initialize session if needed
        default_session = "default"
        if default_session not in file_contexts:
            file_contexts[default_session] = {}
        
        file_contexts[default_session][file_id] = extracted_text
        
        # Optionally save to disk for persistence
        context_file = CONTEXT_DIR / f"{file_id}.txt"
        with open(context_file, "w", encoding="utf-8") as f:
            f.write(extracted_text)
        
        return FileUploadResponse(
            file_id=file_id,
            filename=file.filename,
            size=len(content),
            status="success",
            extracted_chars=len(extracted_text),
            extracted_text=extracted_text[:1000],
            keyword_analysis=keyword_analysis,
            document_analysis=document_analysis
        )
    
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"File processing failed: {str(e)}"
        )


async def generate_streaming_response(session_id: str, messages: List[Dict]):
    """
    Generator function for streaming LLM responses.
    Calls local llama.cpp API and yields tokens in real-time.
    """
    try:
        async with httpx.AsyncClient(timeout=300.0) as client:
            payload = {
                "model": "local-model",
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": 2048,
                "stream": True,
            }
            
            async with client.stream(
                "POST",
                LOCAL_LLM_ENDPOINT,
                json=payload
            ) as response:
                if response.status_code != 200:
                    yield f"data: {{\"error\": \"LLM API returned {response.status_code}\"}}\n\n"
                    return
                
                async for line in response.aiter_lines():
                    if line.startswith("data: "):
                        data_str = line[6:]
                        if data_str.strip() == "[DONE]":
                            yield f"data: {{\"done\": true}}\n\n"
                            break
                        try:
                            chunk = json.loads(data_str)
                            if "choices" in chunk and len(chunk["choices"]) > 0:
                                delta = chunk["choices"][0].get("delta", {})
                                if "content" in delta:
                                    content = delta["content"]
                                    # Send as SSE event
                                    yield f"data: {{\"token\": {json.dumps(content)}}}\n\n"
                        except json.JSONDecodeError:
                            pass
    
    except Exception as e:
        yield f"data: {{\"error\": \"{str(e)}\"}}\n\n"


@app.post("/chat")
async def chat_stream(request: ChatRequest):
    """
    Streaming chat endpoint.
    - Accepts conversation context
    - Injects file context if file_ids provided
    - Streams response from local LLM via SSE
    """
    session_id = request.session_id
    user_message = request.message
    file_ids = request.file_ids or []
    
    # Initialize session if needed
    if session_id not in conversation_history:
        conversation_history[session_id] = []
    if session_id not in file_contexts:
        file_contexts[session_id] = {}
    
    # Build file context
    file_context = ""
    if file_ids:
        for file_id in file_ids:
            if file_id in file_contexts.get(session_id, {}):
                file_context += f"\n[File Context: {file_id}]\n"
                file_context += file_contexts[session_id][file_id]
                file_context += "\n"
    
    # Build system prompt with file context
    system_prompt = "You are a helpful assistant."
    if file_context:
        system_prompt += f"\n\nYou have access to the following documents:\n{file_context}"
    
    # Prepare messages for LLM
    messages = [
        {"role": "system", "content": system_prompt}
    ]
    
    # Add conversation history
    for msg in conversation_history[session_id]:
        messages.append(msg)
    
    # Add current user message
    messages.append({"role": "user", "content": user_message})
    
    # Store user message in history
    conversation_history[session_id].append({
        "role": "user",
        "content": user_message
    })
    
    # Return streaming response
    async def event_generator():
        """Generate SSE events"""
        full_response = ""
        async for chunk in generate_streaming_response(session_id, messages):
            full_response += chunk
            yield chunk
        
        # After streaming completes, extract and store assistant message
        try:
            # Parse the response to extract final message
            if full_response:
                conversation_history[session_id].append({
                    "role": "assistant",
                    "content": "Response streamed (see client logs)"
                })
        except:
            pass
    
    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream"
    )


@app.get("/sessions/{session_id}/history")
async def get_history(session_id: str):
    """Retrieve conversation history for a session"""
    return {
        "session_id": session_id,
        "history": conversation_history.get(session_id, []),
        "file_ids": list(file_contexts.get(session_id, {}).keys())
    }


@app.delete("/sessions/{session_id}")
async def clear_session(session_id: str):
    """Clear conversation history and file contexts for a session"""
    if session_id in conversation_history:
        del conversation_history[session_id]
    if session_id in file_contexts:
        del file_contexts[session_id]
    return {"status": "cleared", "session_id": session_id}


@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "upload_dir": str(UPLOAD_DIR),
        "context_dir": str(CONTEXT_DIR)
    }


@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "name": "Offline File Chat API",
        "version": "1.0.0",
        "endpoints": {
            "upload": "POST /upload",
            "chat": "POST /chat (streaming)",
            "history": "GET /sessions/{session_id}/history",
            "clear": "DELETE /sessions/{session_id}",
            "health": "GET /health"
        }
    }


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    uvicorn.run(
        app,
        host="0.0.0.0",
        port=8000,
        log_level="info"
    )
